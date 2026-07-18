"""
Free AI Text-to-Speech web app.

Uses Microsoft Edge's built-in "Read Aloud" neural voices via the open-source
edge-tts library (https://github.com/rany2/edge-tts). No API key, no signup,
no billing - it talks to the same free service the Edge browser uses.

Run locally:
    pip install -r requirements.txt
    python app.py
Then open http://127.0.0.1:5000 in your browser.
"""

import asyncio
import io
import os
import re
import shutil
import tempfile
import threading
import time
import uuid
import zipfile
from datetime import datetime

import edge_tts
from docx import Document
from flask import Flask, jsonify, render_template, request, send_file

app = Flask(__name__)

# Max characters sent to the TTS service in a single request. Edge's service
# can choke on extremely long single requests, so long text gets split into
# chunks (on paragraph/sentence boundaries) and stitched back together.
CHUNK_SIZE = 3000

# A curated, friendly subset of voices instead of the full 300+ list.
# (Full list can be fetched from edge_tts.list_voices() if you want more.)
CURATED_VOICES = [
    {"id": "en-US-AndrewNeural", "label": "Andrew (US, Male)"},
    {"id": "en-US-AriaNeural", "label": "Aria (US, Female)"},
    {"id": "en-US-GuyNeural", "label": "Guy (US, Male)"},
    {"id": "en-US-JennyNeural", "label": "Jenny (US, Female)"},
    {"id": "en-US-EmmaNeural", "label": "Emma (US, Female)"},
    {"id": "en-GB-RyanNeural", "label": "Ryan (UK, Male)"},
    {"id": "en-GB-SoniaNeural", "label": "Sonia (UK, Female)"},
    {"id": "en-AU-NatashaNeural", "label": "Natasha (Australia, Female)"},
    {"id": "en-AU-WilliamNeural", "label": "William (Australia, Male)"},
    {"id": "en-IN-NeerjaNeural", "label": "Neerja (India, Female)"},
    {"id": "es-ES-AlvaroNeural", "label": "Álvaro (Spanish, Male)"},
    {"id": "es-MX-DaliaNeural", "label": "Dalia (Spanish/Mexico, Female)"},
    {"id": "fr-FR-HenriNeural", "label": "Henri (French, Male)"},
    {"id": "fr-FR-DeniseNeural", "label": "Denise (French, Female)"},
    {"id": "de-DE-ConradNeural", "label": "Conrad (German, Male)"},
    {"id": "de-DE-KatjaNeural", "label": "Katja (German, Female)"},
    {"id": "it-IT-DiegoNeural", "label": "Diego (Italian, Male)"},
    {"id": "pt-BR-AntonioNeural", "label": "Antônio (Portuguese/Brazil, Male)"},
    {"id": "pt-BR-FranciscaNeural", "label": "Francisca (Portuguese/Brazil, Female)"},
    {"id": "ja-JP-KeitaNeural", "label": "Keita (Japanese, Male)"},
    {"id": "ja-JP-NanamiNeural", "label": "Nanami (Japanese, Female)"},
    {"id": "zh-CN-YunxiNeural", "label": "Yunxi (Chinese, Male)"},
    {"id": "zh-CN-XiaoxiaoNeural", "label": "Xiaoxiao (Chinese, Female)"},
]

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def split_text(text, max_len=CHUNK_SIZE):
    """Split text into chunks under max_len, breaking on paragraph/sentence
    boundaries where possible so audio doesn't cut off mid-sentence."""
    text = text.strip()
    if len(text) <= max_len:
        return [text]

    # First split on paragraphs, then sentences if a paragraph is still too long.
    paragraphs = re.split(r"\n\s*\n", text)
    chunks = []
    current = ""

    def flush():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 <= max_len:
            current += ("\n\n" if current else "") + para
            continue
        flush()
        if len(para) <= max_len:
            current = para
            continue
        # Paragraph itself is too long: split on sentences.
        sentences = re.split(r"(?<=[.!?])\s+", para)
        for sent in sentences:
            if len(current) + len(sent) + 1 <= max_len:
                current += (" " if current else "") + sent
            else:
                flush()
                current = sent
    flush()
    return chunks


async def synthesize_chunk(text, voice, rate, pitch, out_path):
    rate_str = f"{'+' if rate >= 0 else ''}{rate}%"
    pitch_str = f"{'+' if pitch >= 0 else ''}{pitch}Hz"
    communicate = edge_tts.Communicate(text, voice, rate=rate_str, pitch=pitch_str)
    await communicate.save(out_path)


async def synthesize_chunk_with_retry(text, voice, rate, pitch, out_path, attempts=2):
    """Try synthesizing a single chunk, retrying once on transient failures
    (e.g. a dropped connection to the speech service) before giving up on
    just this chunk."""
    last_error = None
    for attempt in range(1, attempts + 1):
        try:
            await synthesize_chunk(text, voice, rate, pitch, out_path)
            return True, None
        except Exception as exc:  # noqa: BLE001
            last_error = exc
            if attempt < attempts:
                await asyncio.sleep(1.5)
    return False, str(last_error)


def concat_mp3s(paths, out_path):
    """Concatenate MP3 files by simple byte concatenation.
    This works reliably for playback in virtually all players/browsers
    since consecutive MPEG audio frames don't require a shared container."""
    with open(out_path, "wb") as outfile:
        for p in paths:
            with open(p, "rb") as infile:
                outfile.write(infile.read())


def extract_text_from_txt(file_storage):
    raw = file_storage.read()
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_text_from_docx(file_storage):
    document = Document(file_storage)
    parts = []
    for para in document.paragraphs:
        parts.append(para.text)
    # Also pull text out of any tables, since paragraphs alone skip them.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n\n".join(p for p in parts if p.strip())


@app.route("/")
def index():
    return render_template("index.html", voices=CURATED_VOICES)


@app.route("/api/extract-text", methods=["POST"])
def extract_text():
    if "file" not in request.files:
        return jsonify({"error": "No file was uploaded."}), 400

    file_storage = request.files["file"]
    filename = file_storage.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".txt":
            text = extract_text_from_txt(file_storage)
        elif ext == ".docx":
            text = extract_text_from_docx(file_storage)
        else:
            return jsonify({"error": "Only .txt and .docx files are supported."}), 400
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": f"Couldn't read that file: {exc}"}), 400

    text = text.strip()
    if not text:
        return jsonify({"error": "No readable text was found in that file."}), 400

    return jsonify({"text": text})



# In-memory job store. A job represents one "generate" request; it runs in a
# background thread so the HTTP request that starts it can return instantly,
# and the page polls /api/generate/status/<job_id> for progress. This avoids
# ever holding a single HTTP request open for the full duration of a long
# synthesis job, which is what was hitting the hosting platform's own proxy
# timeout (a limit separate from anything configurable in this app) and
# causing it to return an HTML error page instead of JSON.
JOBS = {}
JOBS_LOCK = threading.Lock()
JOB_MAX_AGE_SECONDS = 3600  # stale jobs are dropped from memory after an hour


def _prune_old_jobs():
    cutoff = time.time() - JOB_MAX_AGE_SECONDS
    stale = [jid for jid, j in JOBS.items() if j.get("created", 0) < cutoff]
    for jid in stale:
        JOBS.pop(jid, None)


def run_generation_job(job_id, chunks, voice, rate, pitch, mode):
    work_dir = tempfile.mkdtemp(prefix=f"tts_{job_id}_")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ordered_paths = [None] * len(chunks)

    async def run():
        for i, chunk in enumerate(chunks):
            out_path = os.path.join(work_dir, f"part_{i:03d}.mp3")
            ok, error = await synthesize_chunk_with_retry(chunk, voice, rate, pitch, out_path)
            if ok:
                part_name = f"speech_{job_id}_part{i + 1:02d}.mp3"
                part_final_path = os.path.join(OUTPUT_DIR, part_name)
                shutil.move(out_path, part_final_path)
                ordered_paths[i] = part_final_path
                with JOBS_LOCK:
                    JOBS[job_id]["parts"].append({
                        "index": i + 1,
                        "filename": part_name,
                        "download_url": f"/download/{part_name}",
                        "play_url": f"/play/{part_name}",
                    })
            else:
                with JOBS_LOCK:
                    JOBS[job_id]["failed"].append({"index": i + 1, "error": error})

    try:
        asyncio.run(run())
    except Exception as exc:  # noqa: BLE001
        with JOBS_LOCK:
            JOBS[job_id]["error"] = f"Speech generation failed: {exc}"
            JOBS[job_id]["status"] = "done"
        _cleanup_dir(work_dir)
        return

    succeeded_paths = [p for p in ordered_paths if p]

    with JOBS_LOCK:
        if not succeeded_paths:
            first_error = JOBS[job_id]["failed"][0]["error"] if JOBS[job_id]["failed"] else "Unknown error."
            JOBS[job_id]["error"] = (
                f"Speech generation failed for all {len(chunks)} part(s). First error: {first_error}"
            )
            JOBS[job_id]["status"] = "done"
        else:
            final = None
            if mode == "zip" and len(chunks) > 1:
                archive_name = f"speech_{stamp}_{job_id}.zip"
                archive_path = os.path.join(OUTPUT_DIR, archive_name)
                with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    for idx, p in enumerate(ordered_paths):
                        if p:
                            zf.write(p, arcname=f"part_{idx + 1:02d}.mp3")
                final = {"download_url": f"/download/{archive_name}", "filename": archive_name}
            elif mode != "parts":
                if len(succeeded_paths) == 1:
                    # Only one chunk total: reuse its file directly as the final
                    # output rather than creating a needless duplicate copy.
                    only_part = JOBS[job_id]["parts"][0]
                    final = {
                        "download_url": only_part["download_url"],
                        "play_url": only_part["play_url"],
                        "filename": only_part["filename"],
                    }
                else:
                    final_name = f"speech_{stamp}_{job_id}.mp3"
                    final_path = os.path.join(OUTPUT_DIR, final_name)
                    concat_mp3s(succeeded_paths, final_path)
                    final = {
                        "download_url": f"/download/{final_name}",
                        "play_url": f"/play/{final_name}",
                        "filename": final_name,
                    }
            JOBS[job_id]["final"] = final
            JOBS[job_id]["status"] = "done"

    _cleanup_dir(work_dir)


def _cleanup_dir(work_dir):
    try:
        for f in os.listdir(work_dir):
            os.remove(os.path.join(work_dir, f))
        os.rmdir(work_dir)
    except OSError:
        pass


@app.route("/api/generate/start", methods=["POST"])
def generate_start():
    data = request.get_json(force=True)
    text = (data.get("text") or "").strip()
    voice = data.get("voice") or "en-US-EmmaNeural"
    rate = int(data.get("rate", 0))
    pitch = int(data.get("pitch", 0))
    mode = data.get("mode", "single")
    if mode not in ("single", "zip", "parts"):
        mode = "single"

    if not text:
        return jsonify({"error": "No text provided."}), 400

    voice_ids = {v["id"] for v in CURATED_VOICES}
    if voice not in voice_ids:
        return jsonify({"error": "Unknown voice."}), 400

    rate = max(-50, min(50, rate))
    pitch = max(-50, min(50, pitch))

    chunks = split_text(text)
    job_id = uuid.uuid4().hex[:12]

    with JOBS_LOCK:
        _prune_old_jobs()
        JOBS[job_id] = {
            "status": "running",
            "mode": mode,
            "total_chunks": len(chunks),
            "parts": [],
            "failed": [],
            "final": None,
            "error": None,
            "created": time.time(),
        }

    thread = threading.Thread(
        target=run_generation_job,
        args=(job_id, chunks, voice, rate, pitch, mode),
        daemon=True,
    )
    thread.start()

    return jsonify({"job_id": job_id, "total_chunks": len(chunks), "mode": mode})


@app.route("/api/generate/status/<job_id>")
def generate_status(job_id):
    with JOBS_LOCK:
        job = JOBS.get(job_id)
        if not job:
            return jsonify({"error": "Unknown or expired job."}), 404
        return jsonify({
            "status": job["status"],
            "mode": job["mode"],
            "total_chunks": job["total_chunks"],
            "parts": list(job["parts"]),
            "failed": list(job["failed"]),
            "final": job["final"],
            "error": job["error"],
        })


@app.route("/download/<path:filename>")
def download(filename):
    safe_name = os.path.basename(filename)
    path = os.path.join(OUTPUT_DIR, safe_name)
    if not os.path.isfile(path):
        return "File not found.", 404
    return send_file(path, as_attachment=True, download_name=safe_name)


@app.route("/play/<path:filename>")
def play(filename):
    """Same as /download but without forcing a download, so it can be used
    directly as an <audio> element's source for immediate playback."""
    safe_name = os.path.basename(filename)
    path = os.path.join(OUTPUT_DIR, safe_name)
    if not os.path.isfile(path):
        return "File not found.", 404
    return send_file(path, as_attachment=False, mimetype="audio/mpeg")


if __name__ == "__main__":
    # Render (and most hosting platforms) provide the port to use via the
    # PORT environment variable, and expect the app to listen on 0.0.0.0
    # (all network interfaces) rather than just 127.0.0.1 (localhost only).
    # Locally, this still defaults to the familiar http://127.0.0.1:5000.
    port = int(os.environ.get("PORT", 5000))
    debug_mode = os.environ.get("FLASK_DEBUG", "0") == "1"
    app.run(host="0.0.0.0", port=port, debug=debug_mode)
